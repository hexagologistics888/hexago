"""
Generate a professional handover .docx from the markdown source + latest project state.
Run: python build_handover_docx.py
"""
import os
from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

BRAND_ORANGE = RGBColor(0xFF, 0x6B, 0x35)
BRAND_NAVY   = RGBColor(0x00, 0x28, 0x55)
BRAND_DARK   = RGBColor(0x0A, 0x16, 0x28)
TEXT_BODY    = RGBColor(0x1A, 0x1A, 0x2E)
TEXT_MUTED   = RGBColor(0x64, 0x74, 0x8B)
TABLE_HEADER_BG = "002855"
TABLE_ALT_BG    = "F8FAFC"
ACCENT_BG       = "FFF4EF"

OUT = r"c:\Hexago Website Trail\Hexago Website Trail\Hexago_Website_Handover_Documentation.docx"


def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def set_cell_borders(cell, color="E2E8F0", size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), size)
        b.set(qn('w:color'), color)
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def set_paragraph_bg(paragraph, hex_color):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hex_color)
    p_pr.append(shd)


def add_horizontal_line(paragraph, color="FF6B35", size="18"):
    p_pr = paragraph._p.get_or_add_pPr()
    bottom = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single')
    b.set(qn('w:sz'), size)
    b.set(qn('w:space'), '1')
    b.set(qn('w:color'), color)
    bottom.append(b)
    p_pr.append(bottom)


def add_heading(doc, text, level=1, color=None, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 0:
        run.font.size = Pt(size or 28)
        run.font.color.rgb = color or RGBColor(0xFF, 0xFF, 0xFF)
    elif level == 1:
        run.font.size = Pt(size or 20)
        run.font.color.rgb = color or BRAND_NAVY
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(8)
        add_horizontal_line(p, color="FF6B35", size="12")
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = color or BRAND_NAVY
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
    elif level == 3:
        run.font.size = Pt(12)
        run.font.color.rgb = color or BRAND_ORANGE
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
    run.font.name = 'Segoe UI'
    return p


def add_body(doc, text, bold=False, color=None, size=11, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Segoe UI'
    run.font.color.rgb = color or TEXT_BODY
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.35
    return p


def add_bullet(doc, text, indent=0, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 + indent * 0.25)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.3
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(11)
        r.font.name = 'Segoe UI'
        r.font.color.rgb = TEXT_BODY
        r2 = p.add_run(text)
        r2.font.size = Pt(11)
        r2.font.name = 'Segoe UI'
        r2.font.color.rgb = TEXT_BODY
    else:
        r = p.add_run(text)
        r.font.size = Pt(11)
        r.font.name = 'Segoe UI'
        r.font.color.rgb = TEXT_BODY
    return p


def add_callout(doc, title, body, color_bg="FFF4EF", color_border="FF6B35"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    set_cell_bg(cell, color_bg)
    set_cell_borders(cell, color=color_border, size="8")
    cell.width = Inches(6.5)
    p_title = cell.paragraphs[0]
    r_title = p_title.add_run(title)
    r_title.bold = True
    r_title.font.color.rgb = BRAND_ORANGE
    r_title.font.size = Pt(11)
    r_title.font.name = 'Segoe UI'
    p_body = cell.add_paragraph()
    r_body = p_body.add_run(body)
    r_body.font.color.rgb = TEXT_BODY
    r_body.font.size = Pt(10.5)
    r_body.font.name = 'Segoe UI'
    p_body.paragraph_format.line_spacing = 1.4
    doc.add_paragraph()


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    # header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_bg(cell, TABLE_HEADER_BG)
        set_cell_borders(cell, color="002855", size="6")
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(10.5)
        r.font.name = 'Segoe UI'
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            if ri % 2 == 1:
                set_cell_bg(cell, TABLE_ALT_BG)
            set_cell_borders(cell, color="E2E8F0", size="4")
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(str(val))
            r.font.size = Pt(10)
            r.font.name = 'Segoe UI'
            r.font.color.rgb = TEXT_BODY
    doc.add_paragraph()
    return table


def add_cover(doc):
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("HEXA GO")
    r.bold = True
    r.font.size = Pt(48)
    r.font.color.rgb = BRAND_NAVY
    r.font.name = 'Segoe UI'
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("LOGISTICS AND TRANSPORT")
    r2.font.size = Pt(12)
    r2.font.color.rgb = BRAND_ORANGE
    r2.font.name = 'Segoe UI'
    r2.bold = True

    p_div = doc.add_paragraph()
    p_div.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_horizontal_line(p_div, color="FF6B35", size="18")

    doc.add_paragraph()
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("Website Project Handover")
    r3.bold = True
    r3.font.size = Pt(26)
    r3.font.color.rgb = BRAND_DARK
    r3.font.name = 'Segoe UI'

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run("Detailed Technical & Operational Documentation")
    r4.font.size = Pt(13)
    r4.font.color.rgb = TEXT_MUTED
    r4.font.name = 'Segoe UI'
    r4.italic = True

    for _ in range(6):
        doc.add_paragraph()

    meta_tbl = doc.add_table(rows=5, cols=2)
    meta_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta = [
        ("Project", "Hexa Go Logistics and Transport — Corporate Website"),
        ("Live URL", "https://hexagologistics.com"),
        ("Handover Date", date.today().strftime("%d %B %Y")),
        ("Prepared For", "Hexa Go Logistics and Transport Team"),
        ("Prepared By", "Development Team"),
    ]
    for i, (k, v) in enumerate(meta):
        c1 = meta_tbl.rows[i].cells[0]
        c2 = meta_tbl.rows[i].cells[1]
        c1.width = Inches(1.8)
        c2.width = Inches(4.2)
        set_cell_bg(c1, TABLE_HEADER_BG)
        set_cell_bg(c2, "FFFFFF")
        set_cell_borders(c1, color="002855", size="4")
        set_cell_borders(c2, color="E2E8F0", size="4")
        pk = c1.paragraphs[0]
        pk.paragraph_format.space_before = Pt(4)
        pk.paragraph_format.space_after = Pt(4)
        rk = pk.add_run(k)
        rk.bold = True
        rk.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        rk.font.size = Pt(10.5)
        rk.font.name = 'Segoe UI'
        pv = c2.paragraphs[0]
        pv.paragraph_format.space_before = Pt(4)
        pv.paragraph_format.space_after = Pt(4)
        rv = pv.add_run(v)
        rv.font.color.rgb = TEXT_BODY
        rv.font.size = Pt(10.5)
        rv.font.name = 'Segoe UI'

    doc.add_page_break()


def add_toc(doc):
    add_heading(doc, "Table of Contents", level=1)
    toc = [
        "1. Executive Summary",
        "2. Project Overview",
        "3. Performance Achievements",
        "4. Technology Stack",
        "5. File Structure & Repository",
        "6. Website Pages",
        "7. Hero Section & Homepage Design",
        "8. Deployment (Vercel)",
        "9. Contact Forms & Lead Capture",
        "10. Careers Application System",
        "11. Email System (Google Apps Script)",
        "12. SEO Configuration",
        "13. Service Worker & PWA",
        "14. Security Headers & CSP",
        "15. Design System & Brand Guidelines",
        "16. How to Make Common Updates",
        "17. Third-Party Services & Accounts",
        "18. Maintenance Checklist",
        "19. Troubleshooting Guide",
        "20. Future Enhancement Recommendations",
        "21. Support & Contacts",
    ]
    for item in toc:
        p = doc.add_paragraph()
        r = p.add_run(item)
        r.font.size = Pt(11.5)
        r.font.color.rgb = TEXT_BODY
        r.font.name = 'Segoe UI'
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.35
    doc.add_page_break()


def section_executive_summary(doc):
    add_heading(doc, "1. Executive Summary", level=1)
    add_body(doc,
        "This document is the complete handover package for the Hexa Go Logistics and Transport corporate website. "
        "It covers architecture, deployment, email automation, SEO, performance, maintenance, and operations."
    )
    add_heading(doc, "Project at a Glance", level=2)
    rows = [
        ("Live URL", "https://hexagologistics.com"),
        ("Hosting", "Vercel (static deployment, auto-deploy on git push)"),
        ("Repository", "github.com/hexagologistics888/hexago (main branch)"),
        ("Primary Domain", "hexagologistics.com (www → apex redirect enforced)"),
        ("CDN", "Vercel Edge Network (global)"),
        ("SSL", "Auto-renewed by Vercel (Let's Encrypt)"),
        ("Pages", "12 HTML pages (homepage, 4 service pages, about, contact, careers, privacy, blog, 404)"),
        ("Email Automation", "Google Apps Script (contact form + careers application)"),
        ("Mobile Performance", "95 / 100 (Lighthouse, April 2026)"),
        ("Desktop Performance", "97 / 100 (Lighthouse, April 2026)"),
        ("Accessibility", "100 / 100 (mobile), 95 / 100 (desktop)"),
        ("Best Practices", "100 / 100"),
        ("SEO", "100 / 100"),
    ]
    add_table(doc, ["Attribute", "Value"], rows, col_widths=[1.9, 4.6])

    add_heading(doc, "Key Capabilities Delivered", level=2)
    caps = [
        "Fully responsive, mobile-first design tested across iOS, Android, and desktop browsers",
        "Three-slide animated hero carousel with numbered 01/02/03 progress indicators",
        "Contact form + quote modal with dual notification (email + WhatsApp)",
        "Careers application with automated two-stage email flow to applicants",
        "Progressive Web App (PWA) — installable on mobile, works offline via Service Worker",
        "Enterprise-grade SEO: 8+ structured data schemas, sitemap, robots.txt, AI search readiness",
        "Security hardened: HSTS, CSP, X-Frame-Options, COOP, Referrer-Policy, Permissions-Policy",
        "Performance optimized: critical CSS inlined, deferred loading, WebP images, self-hosted fonts, conditional GSAP (desktop only)",
    ]
    for c in caps:
        add_bullet(doc, c)
    doc.add_page_break()


def section_project_overview(doc):
    add_heading(doc, "2. Project Overview", level=1)
    add_body(doc,
        "Hexa Go Logistics and Transport is a pan-India logistics and transportation company based in Aluva, Kerala. "
        "The website is the primary digital front for the business and serves four main purposes: (1) brand presence and credibility, "
        "(2) service information for FTL, PTL, 3PL Warehousing, and 4PL Supply Chain, (3) lead capture for business enquiries, "
        "and (4) careers/recruitment."
    )
    add_heading(doc, "Target Audience", level=2)
    aud = [
        ("Enterprise shippers", "Manufacturing, FMCG, and e-commerce companies requiring FTL/PTL freight or 3PL warehousing"),
        ("SMB customers", "Regional businesses needing quote-based logistics for specific routes or one-off shipments"),
        ("Job applicants", "Professionals seeking logistics-industry careers (drivers, warehouse staff, ops, sales, admin)"),
        ("Partners & vendors", "Trucking owner-operators, warehouse providers, technology partners"),
    ]
    add_table(doc, ["Audience", "Description"], aud, col_widths=[1.8, 4.7])

    add_heading(doc, "Core Services Represented", level=2)
    svc = [
        ("FTL (Full Truck Load)", "Dedicated-truck freight for large shipments across India"),
        ("PTL (Partial Truck Load)", "Shared-truck freight with 'Coming Soon' badge — launching phase"),
        ("3PL Warehousing", "Warehouse management, inventory, pick-pack-dispatch"),
        ("4PL Supply Chain", "End-to-end supply chain orchestration as integration partner"),
        ("Route Planning", "Optimised multi-stop route design for client operations"),
        ("Technology Integration", "GPS tracking, ERP integrations, custom logistics tech"),
    ]
    add_table(doc, ["Service", "Description"], svc, col_widths=[2.0, 4.5])

    add_callout(doc,
        "BRAND NAMING NOTE",
        "The official brand name is 'Hexa go' (with a space). 'Hexa-go' and 'Hexago' are accepted variants used in "
        "URL slugs and some legacy content. When writing new content, prefer 'Hexa go'. All three variants are listed in "
        "SEO metadata (title, meta description, keywords, hidden span) for search-engine coverage."
    )
    doc.add_page_break()


def section_performance(doc):
    add_heading(doc, "3. Performance Achievements", level=1)
    add_body(doc,
        "The site underwent an aggressive performance overhaul in April 2026. The optimisations below represent the "
        "final state at handover."
    )
    add_heading(doc, "Lighthouse Scores (April 2026)", level=2)
    perf = [
        ("Performance", "95", "97"),
        ("Accessibility", "100", "95"),
        ("Best Practices", "100", "100"),
        ("SEO", "100", "100"),
        ("First Contentful Paint", "1.0 s", "0.5 s"),
        ("Largest Contentful Paint", "2.9 s", "0.7 s"),
        ("Total Blocking Time", "20 ms", "0 ms"),
        ("Cumulative Layout Shift", "0", "0.097"),
        ("Speed Index", "2.3 s", "0.8 s"),
    ]
    add_table(doc, ["Metric", "Mobile", "Desktop"], perf, col_widths=[2.5, 1.7, 2.3])

    add_heading(doc, "Key Optimisations Applied", level=2)
    opts = [
        ("Render-blocking CSS eliminated", "styles.css and enhancements.css use the media='print' onload-swap pattern; critical above-the-fold CSS (~6KB) is inlined in <head>. Saved ~290ms to FCP."),
        ("Bootstrap CSS removed from critical path", "Full Bootstrap CSS (27KB) replaced with a ~3KB inline subset containing only modal, form-control, form-select, and utility classes actually used."),
        ("Bootstrap JS lazy-loaded", "Bundle only downloads when a user clicks a [data-bs-toggle='modal'] trigger. Most visitors never fetch it."),
        ("GSAP gated to desktop only", "GSAP + ScrollTrigger (~110KB) are dynamically injected only when matchMedia(min-width:769px) matches. Mobile never downloads or parses them."),
        ("Mobile reveals via IntersectionObserver", "Replaces all ScrollTrigger logic on mobile with a single native IntersectionObserver. Zero forced reflows, zero layout reads, offloaded to compositor thread."),
        ("Image format & size", "All hero images converted JPG → WebP (~272KB saved). Service images resized to match displayed dimensions (768→720 wide). Total ~245KB saved on mobile initial load."),
        ("JS minified", "enhancements.js minified from 17.4KB → 6.8KB. Source preserved as enhancements.src.js for future edits."),
        ("ScrollTrigger batched (desktop)", "Per-element scrollTriggers consolidated into ScrollTrigger.batch() — 25 triggers reduced to 5."),
        ("Duplicate scroll handler removed", "Inline script at index.html:1117 was duplicating scrollTop logic from script.js — eliminated double layout reads per scroll."),
        ("Forced reflow in enhancements.js fixed", "Preloader visibility check switched from getComputedStyle() to inline-style + class check. Saved ~42ms of layout invalidation."),
        ("Carousel replaced", "Bootstrap carousel (which performed layout reads every 6s) replaced with a 600-byte inline vanilla JS carousel. Zero reflows."),
        ("CLS fixed", "Inline critical CSS now matches final enhancements.css values. Previous mismatch caused a 0.264 layout shift at hero-content."),
    ]
    add_table(doc, ["Optimisation", "Details"], opts, col_widths=[2.3, 4.2])

    add_callout(doc,
        "WHY MOBILE-SPECIFIC PATHS",
        "Mobile devices have weaker CPUs and slower networks. The site uses two code paths: desktop gets rich GSAP animations, "
        "mobile gets pure CSS + IntersectionObserver reveals. This keeps mobile TBT near zero while desktop retains polished motion. "
        "The matchMedia(min-width:769px) check in the head decides which scripts to load."
    )
    doc.add_page_break()


def section_tech_stack(doc):
    add_heading(doc, "4. Technology Stack", level=1)
    rows = [
        ("HTML", "HTML5 with semantic markup (main, section, article, nav, header, footer)"),
        ("CSS", "Vanilla CSS across 4 files: styles.css (main), enhancements.css, service-pages.css, blog.css"),
        ("JavaScript", "Vanilla JS: script.js (main, ~15KB), enhancements.js (minified, 6.8KB), enhancements.src.js (source for edits)"),
        ("Animations (Desktop)", "GSAP 3.12.2 + ScrollTrigger — loaded from cdnjs.cloudflare.com"),
        ("Animations (Mobile)", "Pure CSS transitions + single IntersectionObserver — zero library dependencies"),
        ("CSS Framework", "Bootstrap 5.3.2 — only subset inlined (~3KB) on homepage; other pages load full CSS from CDN"),
        ("Fonts", "Inter (self-hosted in /fonts/inter-latin.woff2) with font-display: swap"),
        ("Images", "WebP (primary) with JPEG/PNG fallbacks via <picture>"),
        ("Hosting", "Vercel (static deployment, global edge CDN)"),
        ("Domain", "hexagologistics.com (apex + www redirect)"),
        ("Backend (Forms)", "Google Apps Script deployed as web apps"),
        ("Backend (Email)", "Gmail API via GmailApp.sendEmail in Apps Script"),
        ("Icons", "Inline SVGs (no icon library)"),
        ("Version Control", "Git + GitHub (github.com/hexagologistics888/hexago)"),
        ("Build tools (optional)", "Terser for JS minification; Sharp CLI for image conversion/resizing"),
    ]
    add_table(doc, ["Component", "Technology"], rows, col_widths=[2.0, 4.5])
    doc.add_page_break()


def section_file_structure(doc):
    add_heading(doc, "5. File Structure & Repository", level=1)
    add_heading(doc, "Repository", level=2)
    add_body(doc, "Host: github.com/hexagologistics888/hexago", bold=True)
    add_body(doc, "Primary branch: main", bold=True)
    add_body(doc, "Deployment trigger: Any push to main auto-deploys via Vercel's Git integration.")

    add_heading(doc, "Top-Level File Map", level=2)
    files = [
        ("HTML Pages (12)", "index.html, about_us.html, service.html, ftl.html, ptl.html, warehousing.html, fourpl.html, contact_us.html, careers.html, privacy_policy.html, blog-ftl-vs-ptl.html, 404.html"),
        ("Stylesheets (4)", "styles.css (main), enhancements.css (UI upgrades + hero), service-pages.css, blog.css"),
        ("Scripts (4)", "script.js (main logic), enhancements.js (minified production), enhancements.src.js (source for edits), sw.js (Service Worker)"),
        ("Config", "vercel.json (headers + redirects), manifest.json / site.webmanifest (PWA), sitemap.xml, robots.txt, llms.txt (AI crawlers)"),
        ("Assets", "images/ (70+ files, WebP + JPEG), fonts/ (Inter woff2 + CSS), videos/ (mp4), emails/ (templates)"),
        ("Favicons", "favicon.ico, favicon-16x16.png, favicon-32x32.png, apple-touch-icon.png, android-chrome-192x192.png, android-chrome-512x512.png"),
        ("Docs", "Hexago_Website_Handover_Documentation.docx (this file), Hexago_Handover_Documentation.md (source)"),
    ]
    add_table(doc, ["Category", "Files"], files, col_widths=[1.8, 4.7])

    add_callout(doc,
        "enhancements.js BUILD NOTE",
        "enhancements.js is the MINIFIED production build (~6.8KB). The readable source is enhancements.src.js. "
        "To edit behavior, modify enhancements.src.js, then re-minify with: "
        "npx --yes terser@5 enhancements.src.js --compress --mangle --output enhancements.js. "
        "All HTML files reference enhancements.js — no template changes needed after rebuild."
    )
    doc.add_page_break()


def section_pages(doc):
    add_heading(doc, "6. Website Pages", level=1)
    pages = [
        ("Homepage", "/", "Hero carousel, stats bar, services grid, sectors, testimonials, blog teasers, CTA banner, footer"),
        ("About Us", "/about_us.html", "Company story, mission, founders (with Person schema for E-E-A-T), FAQ"),
        ("Services", "/service.html", "Overview of all 6 services with gallery and deep-dive CTAs"),
        ("FTL Services", "/ftl.html", "Full Truck Load dedicated landing page"),
        ("PTL Services", "/ptl.html", "Partial Truck Load page (with 'Coming Soon' badge site-wide)"),
        ("Warehousing", "/warehousing.html", "3PL Warehouse Management deep-dive"),
        ("4PL Supply Chain", "/fourpl.html", "Fourth-Party Logistics / supply chain orchestration page"),
        ("Contact Us", "/contact_us.html", "Form + phone + email + embedded Google Map"),
        ("Careers", "/careers.html", "Job openings + application form (separate Apps Script)"),
        ("Privacy Policy", "/privacy_policy.html", "Data collection, cookies, rights"),
        ("Blog: FTL vs PTL", "/blog-ftl-vs-ptl.html", "Educational article driving organic traffic"),
        ("404 Page", "/404.html", "Custom error page with offline-detection banner"),
    ]
    add_table(doc, ["Page", "URL Path", "Purpose"], pages, col_widths=[1.5, 1.7, 3.3])
    doc.add_page_break()


def section_hero(doc):
    add_heading(doc, "7. Hero Section & Homepage Design", level=1)
    add_body(doc,
        "The homepage hero was rebuilt to enterprise-tier composition in April 2026. It is the most visible piece of the site "
        "and represents roughly 70% of the above-the-fold real estate on first load."
    )
    add_heading(doc, "Current Composition", level=2)
    items = [
        ("Kicker eyebrow", "Orange uppercase label with flanking orange dashes — 'HEXA GO LOGISTICS AND TRANSPORT' style"),
        ("Headline H1 / H2", "Two-line format: white main phrase + orange supporting phrase (e.g. 'Freight Built on Trust. / Delivered Across India.')"),
        ("Supporting paragraph", "Benefit-led copy with target keywords (FTL, PTL, 3PL, 4PL, pincodes, states)"),
        ("Single CTA", "Centered ghost button with rightward arrow → (e.g. 'Request a Quote')"),
        ("Numbered indicators", "01 / 02 / 03 with thin progress bars that fill orange on active slide (replaces Bootstrap dashes)"),
        ("Image treatment", "WebP hero images with directional gradient overlay (clear top-to-dark bottom) + contrast/saturation filter for vibrance"),
    ]
    for i, (k, v) in enumerate(items):
        add_bullet(doc, v, bold_prefix=f"{k}: ")

    add_heading(doc, "Three Slides", level=2)
    slides = [
        ("Slide 1", "Freight Built on Trust. Delivered Across India.", "Request a Quote", "Brand-level intro + full stack of services"),
        ("Slide 2", "On Time. Every Time.", "Get a Quote", "Reliability & coverage (7+ states, 1,000+ pincodes)"),
        ("Slide 3", "End-to-End Logistics. Uncompromising Standards.", "Talk to Our Team", "Full-chain visibility value proposition"),
    ]
    add_table(doc, ["Slide", "Headline", "CTA", "Theme"], slides, col_widths=[0.9, 2.6, 1.3, 1.7])

    add_heading(doc, "SEO Preservation", level=2)
    seo = [
        "H1 appears only on slide 1 with primary keywords intact",
        "Hidden span preserves 'Hexago' / 'Hexa-go' variants for keyword coverage",
        "Alt text on hero image: 'Hexa go Logistics and Transport fleet of freight trucks across India'",
        "Numbered indicators use semantic <button> with aria-label",
    ]
    for s in seo:
        add_bullet(doc, s)
    doc.add_page_break()


def section_deployment(doc):
    add_heading(doc, "8. Deployment (Vercel)", level=1)
    add_heading(doc, "How Deployment Works", level=2)
    add_body(doc,
        "The site is a static deployment on Vercel linked to the GitHub repository. Every commit to the 'main' branch triggers "
        "an automatic production deployment. Preview deployments are created for every pull request."
    )
    add_heading(doc, "vercel.json Responsibilities", level=2)
    items = [
        ("Security headers", "HSTS with preload, CSP, X-Frame-Options, X-Content-Type-Options, Referrer-Policy, Permissions-Policy, COOP, X-XSS-Protection"),
        ("Redirects", "/about-hexago.html → /about_us.html (301), /index.html → / (301), www.hexagologistics.com → hexagologistics.com (301)"),
        ("404 routing", "All unmatched paths serve /404.html with offline detection"),
        ("Cache-Control", "Long cache for fingerprinted assets, short cache for HTML"),
    ]
    for k, v in items:
        add_bullet(doc, v, bold_prefix=f"{k}: ")

    add_heading(doc, "Redeploy Process", level=2)
    steps = [
        "Make code changes locally and commit to the repository",
        "Push to origin/main — Vercel auto-deploys in 1–2 minutes",
        "Or use Vercel dashboard: Project → Deployments → 'Redeploy' on any commit",
        "Verify deploy at: https://vercel.com/dashboard (requires account access)",
    ]
    for s in steps:
        add_bullet(doc, s)

    add_callout(doc,
        "ACCESS REQUIREMENT",
        "The Vercel account that owns this project MUST remain accessible to the business. Without it, you cannot change "
        "domain settings, environment variables, redirects, or trigger manual redeploys. Request admin access from the "
        "developer if you don't already have it."
    )
    doc.add_page_break()


def section_forms(doc):
    add_heading(doc, "9. Contact Forms & Lead Capture", level=1)
    add_heading(doc, "Contact Form Flow", level=2)
    add_body(doc,
        "The contact form on /contact_us.html and the quote modal on / submit to a Google Apps Script web app that: "
        "(1) stores the lead, (2) emails the admin team, and (3) sends a branded auto-reply to the customer."
    )
    rows = [
        ("Frontend", "form#contactForm (contact_us.html), form#ctaModalForm (homepage modal)"),
        ("Form library", "HexagoLeadForms.attach() in script.js handles validation, submission, success animation"),
        ("Backend URL", "https://script.google.com/macros/s/AKfycbxRm8Qv.../exec (Apps Script web app)"),
        ("Fields captured", "name, email, company, phone, service, message"),
        ("Admin notification", "info@hexagologistics.com (plain + HTML premium template)"),
        ("Customer auto-reply", "Branded HTML email with logo, enquiry summary, CTAs"),
        ("WhatsApp notification", "Optional — composes a WhatsApp message to +91 7034298745 with lead details"),
        ("Success UI", "Branded 'Enquiry Received' popup, form reset, temporary green 'Sent!' button state"),
    ]
    add_table(doc, ["Attribute", "Value"], rows, col_widths=[1.8, 4.7])

    add_callout(doc,
        "GOOGLE ACCOUNT OWNERSHIP",
        "The Google account that owns the Apps Script IS the lifeline of all form submissions. If access is lost or the owner "
        "leaves the company, leads stop being saved. Document and share the owning Google account credentials with a senior "
        "internal contact (preferably a company-owned email, not personal)."
    )
    doc.add_page_break()


def section_careers(doc):
    add_heading(doc, "10. Careers Application System", level=1)
    add_body(doc,
        "The careers page uses a dedicated Apps Script (separate from contact form) with a two-stage automated email flow."
    )
    add_heading(doc, "Application Flow", level=2)
    flow = [
        "1. Applicant fills form on /careers.html (name, email, phone, role, message)",
        "2. Form posts to dedicated careers Apps Script web app",
        "3. Script sends admin notification to hr@hexagologistics.com and kichuharishankar@hexagologistics.com",
        "4. Script immediately sends 'Application Received' confirmation to applicant",
        "5. A scheduled trigger fires ~1 minute later, sending a second 'Please send your resume' email to the applicant",
        "6. Script uses a single shared trigger (not one-per-submission) to stay under Google's 20-trigger quota",
        "7. Logo images from hexagologistics.com/images/logo3.png are embedded in both header and footer of all emails",
    ]
    for f in flow:
        add_bullet(doc, f)

    add_heading(doc, "Email Templates", level=2)
    tmpl = [
        ("Admin notification", "Full applicant detail card, message body, Reply-to-applicant CTA"),
        ("Applicant Email 1 (immediate)", "'Application Received' — reference number, review timeline, 3-step process"),
        ("Applicant Email 2 (1-min delay)", "'Please send your resume' — submission instructions, mailto CTA"),
    ]
    add_table(doc, ["Template", "Contents"], tmpl, col_widths=[2.3, 4.2])

    add_callout(doc,
        "DEPLOYMENT REMINDER",
        "After any Apps Script code change, you MUST click 'Deploy → Manage deployments → Edit → Version: New version → Deploy'. "
        "Just saving the code does NOT update the live web app. This is the #1 cause of 'my changes aren't working' issues."
    )
    doc.add_page_break()


def section_email(doc):
    add_heading(doc, "11. Email System (Google Apps Script)", level=1)
    add_body(doc,
        "Two separate Apps Script projects power the website's email automation. Both send from the Google account that owns "
        "the script, using Gmail's daily quota."
    )
    rows = [
        ("Contact form script", "Handles contact_us + quote modal submissions; sends to info@hexagologistics.com"),
        ("Careers script", "Handles careers.html submissions; sends to hr@ + kichuharishankar@hexagologistics.com"),
        ("Email engine", "GmailApp.sendEmail (built into Apps Script, no SMTP config needed)"),
        ("Daily quota (consumer)", "100 recipients/day per Google account"),
        ("Daily quota (Workspace)", "1,500 recipients/day per Google account"),
        ("Logging", "Logger.log() statements in every script; visible via Apps Script → Executions"),
        ("Template CSS", "Inlined in every email (Gmail strips <style> in <head>, so inline style= attributes used on critical elements)"),
        ("Images", "Hosted on hexagologistics.com (logo3.png) — referenced by absolute URL in <img src>"),
    ]
    add_table(doc, ["Element", "Details"], rows, col_widths=[2.0, 4.5])

    add_heading(doc, "Common Email Issues & Fixes", level=2)
    issues = [
        ("Applicants not receiving emails", "Check Apps Script → Executions → most recent run for log output. Common cause: daily quota exceeded, spam folder, or stale deployment"),
        ("Two-step careers flow stops working", "Trigger quota exceeded (20 max). Open Apps Script → Triggers → delete stale 'sendPendingResumeEmails' triggers"),
        ("Images don't render in email", "Gmail blocks .webp from some domains — always use .png for email-embedded logos"),
        ("Emails go to spam", "Add SPF/DKIM for the sending domain; use a consistent From name ('Hexa-go HR Team')"),
    ]
    add_table(doc, ["Symptom", "Fix"], issues, col_widths=[2.3, 4.2])
    doc.add_page_break()


def section_seo(doc):
    add_heading(doc, "12. SEO Configuration", level=1)
    add_heading(doc, "Structured Data (Schema.org JSON-LD)", level=2)
    schemas = [
        ("LocalBusiness + FreightForwarder", "Google Maps, Local Pack, Knowledge Panel"),
        ("Organization", "Knowledge Graph entity for brand queries"),
        ("WebSite with SearchAction", "Sitelinks search box in SERP"),
        ("ItemList (Services)", "Rich service listings"),
        ("Speakable", "Voice assistant readability (Alexa, Google Assistant)"),
        ("Review + AggregateRating", "Star ratings in SERP"),
        ("FAQPage", "FAQ rich results (flagged as Info priority — commercial sites get AI/LLM citation benefit, not Google rich snippets since Aug 2023)"),
        ("Person (×2)", "Founder entities for E-E-A-T signals"),
    ]
    add_table(doc, ["Schema Type", "Purpose"], schemas, col_widths=[2.3, 4.2])

    add_heading(doc, "Technical SEO Checklist", level=2)
    chk = [
        "Title tag, meta description, canonical URL on every page",
        "Open Graph + Twitter Card for social previews",
        "hreflang tags (en-in, en, x-default)",
        "Geo meta tags (region: IN-KL, Kerala, India)",
        "robots meta + robots.txt (allows major crawlers AND AI crawlers: GPTBot, ClaudeBot, PerplexityBot)",
        "sitemap.xml covering all 11 indexable pages",
        "llms.txt file for AI search engines (ChatGPT, Perplexity citations)",
        "Image alt text on every content image",
        "Heading hierarchy: H1 unique per page, H2→H3 sequential (no skipped levels)",
        "Core Web Vitals: INP used (not deprecated FID)",
        "Page speed: 95 mobile / 97 desktop",
    ]
    for c in chk:
        add_bullet(doc, c)

    add_callout(doc,
        "AI SEARCH READY",
        "The site is one of the first logistics companies in India explicitly targeting AI search (ChatGPT web search, "
        "Perplexity, Claude). robots.txt allows GPTBot, ClaudeBot, PerplexityBot, Meta-ExternalAgent, and Amazonbot. "
        "llms.txt provides a structured company summary for AI crawlers. Monitor brand mentions in AI responses monthly."
    )
    doc.add_page_break()


def section_pwa(doc):
    add_heading(doc, "13. Service Worker & PWA", level=1)
    add_heading(doc, "Service Worker (sw.js)", level=2)
    rows = [
        ("Cache version", "hexago-v41 (increment to v42, v43, ... on every major update)"),
        ("HTML strategy", "Network-first with 4s timeout → cache → offline fallback"),
        ("CSS/JS strategy", "Network-first with cache fallback"),
        ("Image/Font strategy", "Cache-first with network fill"),
        ("Pre-cache", "Core HTML pages, stylesheets, main JS, key images"),
        ("Offline fallback", "/404.html with offline-detection banner"),
    ]
    add_table(doc, ["Attribute", "Value"], rows, col_widths=[1.8, 4.7])

    add_heading(doc, "PWA Manifest (manifest.json)", level=2)
    rows2 = [
        ("App name", "Hexa Go Logistics and Transport"),
        ("Short name", "Hexa Go"),
        ("Theme color", "#FF6B35 (brand orange)"),
        ("Background color", "#002855 (brand navy)"),
        ("Display mode", "Standalone (hides browser chrome when installed)"),
        ("Shortcuts", "'Get a Quote' and 'Our Services'"),
        ("Icons", "16×16 through 512×512 (PNG)"),
    ]
    add_table(doc, ["Property", "Value"], rows2, col_widths=[1.8, 4.7])

    add_callout(doc,
        "CACHE VERSION BUMP",
        "Increment CACHE_NAME in sw.js (currently 'hexago-v41') every time you deploy a meaningful update. This forces "
        "returning visitors to get the new version instead of serving stale content from their cache. Without the bump, "
        "users may see the old site for up to 24 hours."
    )
    doc.add_page_break()


def section_security(doc):
    add_heading(doc, "14. Security Headers & CSP", level=1)
    add_body(doc, "Security headers are set in two places for defence-in-depth: Vercel server-side (vercel.json) and HTML meta-tag fallbacks.")

    add_heading(doc, "Headers Applied", level=2)
    hdrs = [
        ("Strict-Transport-Security", "max-age=63072000; includeSubDomains; preload — forces HTTPS"),
        ("Content-Security-Policy", "Restricts script/style/font/image sources to our domain + approved CDNs"),
        ("X-Frame-Options", "SAMEORIGIN — prevents clickjacking via iframe embedding"),
        ("X-Content-Type-Options", "nosniff — prevents MIME type sniffing attacks"),
        ("Referrer-Policy", "strict-origin-when-cross-origin — limits referrer leakage"),
        ("Permissions-Policy", "Disables camera, microphone, geolocation unless explicitly needed"),
        ("Cross-Origin-Opener-Policy", "same-origin — isolates browsing context"),
        ("X-XSS-Protection", "1; mode=block — legacy XSS filter (fallback)"),
    ]
    add_table(doc, ["Header", "Policy"], hdrs, col_widths=[2.3, 4.2])

    add_callout(doc,
        "IMPORTANT — CSP UPDATES",
        "If you add new external services (analytics, chat widgets, new CDNs, embed iframes), you MUST update the "
        "Content-Security-Policy in BOTH vercel.json AND the CSP meta tag in every HTML <head>. Missing CSP updates is "
        "the #1 cause of 'my new integration doesn't work' issues — the browser silently blocks the resource."
    )
    doc.add_page_break()


def section_design_system(doc):
    add_heading(doc, "15. Design System & Brand Guidelines", level=1)
    add_heading(doc, "Brand Colors", level=2)
    colors = [
        ("Primary Orange", "#FF6B35", "CTAs, accents, hero highlights, kicker pills"),
        ("Light Orange", "#FF9A5A", "Hero accent phrase, secondary highlights"),
        ("Deep Red-Orange", "#E8521F", "Button hover / gradient end"),
        ("Dark Navy", "#002855", "Headings, primary brand color"),
        ("Medium Blue", "#004E89", "Secondary blue, stats bar gradient"),
        ("Deep Background", "#0a1628", "Footer, hero overlay tint"),
        ("White", "#FFFFFF", "Text on dark, cards"),
        ("Body Text", "#1A1A1A / #1A1A2E", "Primary body copy"),
        ("Muted", "#495670 / #64748B", "Secondary body, captions"),
        ("Success Green", "#16A34A / #10B981", "Form success states"),
        ("Error Red", "#EF4444", "Form validation errors"),
    ]
    add_table(doc, ["Color", "Hex", "Usage"], colors, col_widths=[1.7, 1.3, 3.5])

    add_heading(doc, "Typography", level=2)
    typ = [
        ("Body", "Inter 400", "16px (1rem), line-height 1.6"),
        ("H1 (Hero)", "Inter 800", "clamp(2rem, 4.2vw, 3.6rem)"),
        ("H2", "Inter 700–800", "Responsive clamp()"),
        ("H3", "Inter 700", "~1.1rem"),
        ("Navigation", "Inter 500", "0.95rem"),
        ("Buttons", "Inter 600–700", "1rem"),
        ("Kicker eyebrow", "Inter 700", "0.72rem, letter-spacing 0.22em, uppercase"),
        ("Micro labels", "Inter 500–600", "0.68–0.72rem, letter-spacing 0.1–0.15em, uppercase"),
    ]
    add_table(doc, ["Element", "Font + Weight", "Size"], typ, col_widths=[1.8, 1.8, 2.9])

    add_heading(doc, "UI Components", level=2)
    comp = [
        ("Cards", "Rounded corners 14px, subtle shadow, fixed image-box height 210px"),
        ("Buttons", "Pill-shaped (50px border-radius), gradient backgrounds, white-ring + orange-ring focus-visible"),
        ("Modal", "Rounded 20px, backdrop-blur overlay, dark navy top-bar accent"),
        ("Form inputs", "Border 1px #ced4da, focus ring rgba(13,110,253,0.25), rounded 0.375rem"),
        ("Hero", "Full-bleed image + directional gradient overlay + numbered 01/02/03 progress indicators"),
        ("Preloader", "Animated hexagonal loader with Hexago branding, hides after 250-800ms based on connection speed"),
    ]
    add_table(doc, ["Component", "Spec"], comp, col_widths=[1.6, 4.9])
    doc.add_page_break()


def section_updates(doc):
    add_heading(doc, "16. How to Make Common Updates", level=1)
    add_heading(doc, "Update Phone Number or Email", level=2)
    for s in [
        "Grep the entire repo for the phone number (e.g. 7034298745) or email address",
        "Replace in ALL HTML files (header nav, footer, contact page, schema JSON-LD)",
        "Update manifest.json, llms.txt, vercel.json (if email is in CSP)",
        "Update the WhatsApp link href (wa.me/91703...)",
        "Commit + push + verify on live site",
    ]:
        add_bullet(doc, s)

    add_heading(doc, "Update Office Address", level=2)
    for s in [
        "Grep repo for 'Vadakumury' or 'Aluva'",
        "Replace in HTML pages and all JSON-LD schema",
        "Update llms.txt and the Google Maps embed URL on contact_us.html",
        "Commit + push",
    ]:
        add_bullet(doc, s)

    add_heading(doc, "Add a New Page", level=2)
    for s in [
        "Copy an existing similar page as a template (e.g. ftl.html → newservice.html)",
        "Update title, meta description, H1, body content, schema markup",
        "Add the page to the <nav> menu in ALL 12 HTML files",
        "Add entry to sitemap.xml with <loc>, <lastmod>, <priority>",
        "Add URL to PRECACHE list in sw.js and increment CACHE_NAME version",
        "Commit + push",
    ]:
        add_bullet(doc, s)

    add_heading(doc, "Add/Update Images", level=2)
    for s in [
        "Export images in both WebP (primary) and PNG/JPEG (fallback)",
        "Target size: WebP < 100KB, JPEG/PNG < 200KB",
        "Place in /images/ directory",
        "In HTML, use <picture> with WebP <source> + JPEG/PNG <img> fallback",
        "Always include width/height attributes to prevent CLS",
        "For hero images, include multiple srcset sizes (768w, 1200w, 1920w)",
    ]:
        add_bullet(doc, s)

    add_heading(doc, "Update Text Content", level=2)
    for s in [
        "Find the relevant HTML file, edit the content directly",
        "If the text appears in schema JSON-LD too (e.g. service names), update there as well",
        "After text changes that affect SEO (titles, meta desc, H1), update sitemap.xml lastmod",
        "Commit + push",
    ]:
        add_bullet(doc, s)

    add_heading(doc, "Add a Blog Post", level=2)
    for s in [
        "Copy blog-ftl-vs-ptl.html as a template",
        "Update title, H1, body, Article schema, dates",
        "Link from homepage blog section",
        "Add to sitemap.xml and sw.js precache",
        "Commit + push",
    ]:
        add_bullet(doc, s)
    doc.add_page_break()


def section_services(doc):
    add_heading(doc, "17. Third-Party Services & Accounts", level=1)
    add_body(doc, "Access to these services is required to maintain the site. Ensure a senior internal contact has credentials for each.")
    rows = [
        ("Vercel", "Website hosting & deployment", "Vercel account login (project admin)", "Critical"),
        ("Domain Registrar", "Ownership of hexagologistics.com", "Registrar login (whoever bought the domain)", "Critical"),
        ("Google Account (Contact Form)", "Apps Script for contact form + Google Sheet data", "Full access to the Google account owning the script", "Critical"),
        ("Google Account (Careers)", "Apps Script for careers application emails", "Full access (may or may not be the same account as above)", "Critical"),
        ("GitHub", "Source code repository", "Admin/push access to github.com/hexagologistics888/hexago", "Critical"),
        ("Google Search Console", "SEO monitoring, sitemap submission", "Verified property access", "Important"),
        ("WhatsApp Business", "+91 7034298745 receives form notifications", "SIM/device with WhatsApp", "Important"),
        ("LinkedIn (company)", "Linked from footer", "Admin access", "Nice-to-have"),
        ("Instagram (company)", "Linked from footer", "Admin access", "Nice-to-have"),
        ("Formspree (optional)", "Alternative email backend (not currently active)", "Formspree account if enabled", "Optional"),
    ]
    add_table(doc, ["Service", "Purpose", "Required", "Priority"], rows, col_widths=[1.5, 2.0, 2.0, 1.0])

    add_heading(doc, "CDN Dependencies (No Account Needed)", level=2)
    cdn = [
        ("Bootstrap 5.3.2", "cdn.jsdelivr.net", "JS only; CSS is inlined"),
        ("GSAP 3.12.2 + ScrollTrigger", "cdnjs.cloudflare.com", "Desktop only"),
        ("Google Fonts", "fonts.googleapis.com", "Fallback only; Inter is self-hosted"),
    ]
    add_table(doc, ["Library", "CDN", "Notes"], cdn, col_widths=[2.0, 2.3, 2.2])
    doc.add_page_break()


def section_maintenance(doc):
    add_heading(doc, "18. Maintenance Checklist", level=1)
    add_heading(doc, "Monthly", level=2)
    for t in [
        "Submit a test lead via contact form — verify Google Sheet + WhatsApp notification arrive",
        "Submit a test careers application — verify both admin and applicant emails land (check spam)",
        "Browse site on mobile and desktop, click through every main page",
        "Check Lighthouse score (mobile + desktop) — flag any regression below 90",
        "Check Google Search Console for crawl errors or indexation drops",
    ]:
        add_bullet(doc, t)

    add_heading(doc, "Quarterly", level=2)
    for t in [
        "Update sitemap.xml <lastmod> dates",
        "Review structured data — update if phone, address, hours, or team changed",
        "Bump sw.js CACHE_NAME (hexago-v41 → hexago-v42) if content updated",
        "Review CSP for any new integrations added",
        "Test offline mode (DevTools → Network → Offline)",
    ]:
        add_bullet(doc, t)

    add_heading(doc, "Annually", level=2)
    for t in [
        "Renew domain registration",
        "Confirm SSL auto-renewal (Vercel handles this automatically)",
        "Update copyright year in every footer",
        "Refresh hero carousel images if product/fleet changed",
        "Review llms.txt and update with any new services or positioning",
        "Audit third-party account access — remove leavers, confirm current staff have it",
    ]:
        add_bullet(doc, t)
    doc.add_page_break()


def section_troubleshoot(doc):
    add_heading(doc, "19. Troubleshooting Guide", level=1)

    def sub(title, steps):
        add_heading(doc, title, level=3)
        for s in steps:
            add_bullet(doc, s)

    sub("Site shows old content after I pushed a change", [
        "Confirm Vercel deployment succeeded (Vercel dashboard → Deployments)",
        "Hard refresh (Ctrl+Shift+R / Cmd+Shift+R)",
        "If still stale, bump sw.js CACHE_NAME and redeploy — users' old service worker is serving cached content",
    ])
    sub("Contact form stopped sending emails", [
        "Open Apps Script → Executions → inspect most recent run's Logger output",
        "Check if Gmail daily quota is exceeded (100 consumer / 1,500 Workspace)",
        "Confirm the web app is still deployed: Deploy → Manage deployments (active version must exist)",
        "Test the Apps Script URL directly in a browser — should return JSON",
    ])
    sub("Careers applicants not getting emails", [
        "Same as contact form — check Executions log",
        "Check Apps Script → Triggers — stale 'sendPendingResumeEmails' triggers accumulate if quota exceeded; delete them",
        "Confirm applicant's email isn't in their spam folder",
    ])
    sub("Images not loading", [
        "Verify the image exists in /images/ and filename matches exactly (case-sensitive on Linux/Vercel)",
        "Check browser DevTools → Network → filter by image — look for 404",
        "Check CSP policy allows the image source (both vercel.json AND HTML meta)",
    ])
    sub("CSS not updating / site looks broken", [
        "Hard refresh first",
        "If the homepage specifically looks unstyled, check inline <style> block in index.html — it carries critical hero/nav CSS",
        "Inspect the styles.css / enhancements.css requests — they use media='print' onload-swap, so if that JS breaks, CSS won't apply",
        "Bump sw.js CACHE_NAME to force cache invalidation",
    ])
    sub("Mobile menu not working", [
        "Check DevTools Console for JS errors — script.js may have failed to parse",
        "Confirm the hamburger has class='mobile-menu-toggle' and the <ul> has id='navLinks'",
        "Check that initMobileMenu() is being called in the DOMContentLoaded handler of script.js",
    ])
    sub("WhatsApp button missing", [
        "Check the .float-call element is present in the page HTML",
        "Verify z-index in enhancements.css is high enough (currently var(--z-whatsapp))",
        "On mobile, check if it's hidden behind the cookie banner or modal",
    ])
    sub("Lighthouse score dropped", [
        "Run Lighthouse again — sometimes a single run gets unlucky network timing",
        "Check the Insights section for new flagged items (new image, new render-blocking script, etc.)",
        "Common regression cause: adding a new script without 'defer' or 'async'",
    ])
    doc.add_page_break()


def section_future(doc):
    add_heading(doc, "20. Future Enhancement Recommendations", level=1)
    add_body(doc, "Suggestions for follow-on work, ordered by business impact.")

    items = [
        ("Blog content expansion", "Currently only 1 post. Aim for 8–12 logistics topics (FTL rates, warehousing costs, 3PL vs 4PL, etc.) to capture organic long-tail traffic"),
        ("Client logos & case studies", "Add a dedicated /case-studies/ page with 3–5 detailed client wins. Strong trust signal for enterprise buyers"),
        ("Live tracking integration", "If Hexa-go offers GPS tracking, add a /track?order=XYZ lookup page — reduces support queries"),
        ("Admin dashboard for leads", "Currently leads land in Google Sheet; a simple Looker Studio or AppSheet dashboard gives sales team better visibility"),
        ("A/B testing on hero CTA", "Test 'Request a Quote' vs 'Get a Free Quote' vs 'Talk to Our Team' — measure conversion uplift"),
        ("FAQ schema on service pages", "Add FAQPage structured data to /ftl.html, /ptl.html, /warehousing.html, /fourpl.html for AI/LLM citations"),
        ("Google Analytics 4 integration", "Add GA4 with Consent Mode v2 — measure traffic sources and conversion paths"),
        ("Email marketing integration", "Connect the careers Apps Script + contact form leads to Mailchimp / Brevo for nurture sequences"),
        ("Video hero background", "Replace static hero image with a short muted MP4 of fleet in motion — stronger brand impression, if load budget allows"),
        ("Multi-language support", "Add Hindi / Malayalam versions using hreflang — significant expansion for Indian market"),
    ]
    add_table(doc, ["Enhancement", "Rationale"], items, col_widths=[1.8, 4.7])
    doc.add_page_break()


def section_support(doc):
    add_heading(doc, "21. Support & Contacts", level=1)
    add_heading(doc, "Primary Site Contact", level=2)
    add_body(doc, "For technical issues, deployment problems, or code questions, contact the developer who built the site.")

    add_heading(doc, "Business Contacts", level=2)
    biz = [
        ("General Enquiries", "info@hexagologistics.com"),
        ("HR / Careers", "hr@hexagologistics.com"),
        ("HR Backup", "kichuharishankar@hexagologistics.com"),
        ("Phone / WhatsApp", "+91 70342 98745"),
        ("Office Hours", "Monday – Saturday, 9:30 AM – 7:30 PM IST"),
        ("Office Address", "Vadakumury Building, Metro Pillar 174, Aluva, Ernakulam, Kerala 683101"),
    ]
    add_table(doc, ["Purpose", "Contact"], biz, col_widths=[1.8, 4.7])

    add_heading(doc, "Useful Links", level=2)
    links = [
        ("Live Website", "https://hexagologistics.com"),
        ("Vercel Dashboard", "https://vercel.com/dashboard"),
        ("Google Search Console", "https://search.google.com/search-console"),
        ("GitHub Repository", "https://github.com/hexagologistics888/hexago"),
        ("Bootstrap 5 Docs", "https://getbootstrap.com/docs/5.3"),
        ("GSAP Docs", "https://gsap.com/docs/v3/"),
        ("Schema.org Validator", "https://validator.schema.org/"),
        ("Google Rich Results Test", "https://search.google.com/test/rich-results"),
        ("PageSpeed Insights", "https://pagespeed.web.dev"),
        ("WCAG 2.2 Reference", "https://www.w3.org/WAI/WCAG22/"),
    ]
    add_table(doc, ["Resource", "URL"], links, col_widths=[2.0, 4.5])

    add_callout(doc,
        "FINAL NOTE",
        "Keep this document updated. Whenever you add a new page, integration, or significant feature, edit the relevant "
        "section so the next person always has accurate information. This document + the source repository are the two "
        "things that keep the site maintainable long-term."
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Document version: 2.0  |  Generated: {date.today().strftime('%d %B %Y')}")
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = TEXT_MUTED
    r.font.name = 'Segoe UI'


def set_default_style(doc):
    style = doc.styles['Normal']
    style.font.name = 'Segoe UI'
    style.font.size = Pt(11)
    style.font.color.rgb = TEXT_BODY


def set_margins(doc):
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)


def main():
    doc = Document()
    set_default_style(doc)
    set_margins(doc)

    add_cover(doc)
    add_toc(doc)
    section_executive_summary(doc)
    section_project_overview(doc)
    section_performance(doc)
    section_tech_stack(doc)
    section_file_structure(doc)
    section_pages(doc)
    section_hero(doc)
    section_deployment(doc)
    section_forms(doc)
    section_careers(doc)
    section_email(doc)
    section_seo(doc)
    section_pwa(doc)
    section_security(doc)
    section_design_system(doc)
    section_updates(doc)
    section_services(doc)
    section_maintenance(doc)
    section_troubleshoot(doc)
    section_future(doc)
    section_support(doc)

    # Write to a new filename so it doesn't clash with an open Word instance
    # The existing file can be replaced by user via rename if needed.
    out_path = OUT
    try:
        doc.save(out_path)
        print(f"WROTE: {out_path}")
    except PermissionError:
        alt = out_path.replace('.docx', '_v2.docx')
        doc.save(alt)
        print(f"Original locked (Word open). Wrote to: {alt}")


if __name__ == '__main__':
    main()
